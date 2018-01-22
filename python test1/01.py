from  docx import  Document         #需要安装docx
from  docx.shared import  Pt
from  docx.oxml.ns import  qn       #需要安装xlmx
from  docx.shared import Inches
from  docx.shared import RGBColor   #导入颜色


#打开文档
#document = Document()
doc_new = Document( "C:\\Users\\Administrator\\Desktop\\报告\\01-电子数据检验委托书-02送检物品清单.docx")
#doc_new = Document("test.docx")
#读取表格


#声明变量获取表格位置元素
table_get1 = doc_new.tables
table_1 = table_get1[0]

'''
#遍历表格
for x in table_1.rows:
    for y in x.cells:
        print(y.text)
'''    

#表格插入内容
#第一个姓名
hdr_cells = table_1.rows[1].cells
hdr_cells[2].text = '张世杰'



#保存word
doc_new.save("C:\\Users\\Administrator\\Desktop\\文档测试\\01-电子数据检验委托书-02送检物品清单.docx")